from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Reusable-Landing-Page-and-Social-Automation-Guide.docx")

INK = "172019"
GREEN = "24352A"
OLIVE = "66743A"
ORANGE = "B85A2B"
IVORY = "F7F5EF"
MUTED = "5F685F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def border_table(table, color="D6DED6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tbl_pr.append(borders)


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(doc, text, style="Normal", size=None, bold=False, color=INK, italic=False, after=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size or 11, bold, color, italic)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + (level * 0.25))
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, {1: 16, 2: 13, 3: 12}[level], True, {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    p.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}[level])
    p.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}[level])
    return p


def code_block(doc, content):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(content)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(INK)
    return table


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    border_table(t)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, item in enumerate(headers):
        cell = hdr.cells[i]
        if widths:
            set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(item), 10, True, GREEN)
    for row in rows:
        cells = t.add_row().cells
        prevent_row_split(t.rows[-1])
        for i, item in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(item), 9.5)
    return t


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        s = doc.styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("OLIVEHORSE FITNESS ACADEMY  |  REUSABLE IMPLEMENTATION GUIDE"), 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_font(footer.add_run("Internal operator reference  •  July 2026"), 8.5, False, MUTED)


def build():
    doc = Document()
    setup(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("REUSABLE IMPLEMENTATION GUIDE"), 11, True, OLIVE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("Landing Pages, Social Creative & Safe Publishing Automation"), 25, True, GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run("A detailed reference and master prompt for the next OliveHorse launch."), 12, False, MUTED)
    table(doc, ["Use this guide for", "Operating principle"], [["Any new landing page, campaign page, blog-to-social workflow, or approved social post.", "Verify business facts, pause for approval, and never publish merely because credentials exist."]], [3.25, 3.25])
    add_text(doc, "Scope", size=12, bold=True, color=GREEN, after=3)
    add_text(doc, "This guide records the implementation standards agreed for OliveHorse Fitness Academy. Replace only bracketed business-specific fields for a new project; retain the safety gates, security controls and visual production rules.", after=12)

    heading(doc, "1. Non-negotiable operating rules")
    for item in [
        "Use only verified business facts. Never invent services, prices, credentials, results, testimonials, locations, contact details, or claims.",
        "Keep React/Vite architecture, npm and ESM unless the project explicitly requires a different stack.",
        "Do not expose, log, commit or paste secrets, OAuth tokens, Blob tokens or upload secrets into source code, state files, screenshots or GitHub Actions output.",
        "Every public landing page must have unique title/description metadata, meaningful HTML and an absolute canonical URL. Draft and temporary paths must stay out of the sitemap.",
        "A topic must be approved before content preparation. A final creative preview must be approved before any real Meta API publish call.",
        "Default to dry-run. AUTO_PUBLISH remains false. Scheduling stays disabled until a verified posting time is configured.",
        "Do not use AI-generated text inside a background image. Render all readable copy as deterministic HTML/SVG text layers.",
    ]: bullet(doc, item)

    heading(doc, "2. Discovery checklist for a new landing page")
    table(doc, ["Area", "Collect and verify before implementation"], [
        ("Business identity", "Official name, address, service area, legal contact method, verified website/canonical domain and approved logo asset."),
        ("Audience & offer", "Who the page serves, programmes/services actually offered, approved CTA and conversion destination."),
        ("Brand system", "Colour hex values, heading/body fonts, logo usage, photo/illustration permissions and accessibility requirements."),
        ("Compliance", "Consent status for identifiable people—especially children—plus any sector-specific advertising restrictions."),
        ("Publishing", "Owner who can approve, preferred social platforms, confirmed time zone, posting time and supported media formats."),
        ("Technical", "Framework, public asset host, Vercel project, env-key names only, analytics/SEO integrations and test commands."),
    ], [1.45, 5.05])
    heading(doc, "3. Master prompt — copy, fill and use", 2)
    code_block(doc, "You are the implementation lead for [BUSINESS NAME]. Work in the existing [FRAMEWORK] repository using [PACKAGE MANAGER] and [MODULE STYLE].\n\nBuild or update [LANDING PAGE / SOCIAL AUTOMATION] without changing unrelated public design. Use only the verified facts below; if information is missing, mark it as verification_required rather than guessing.\n\nVERIFIED BUSINESS PROFILE\n- Canonical website: [HTTPS URL]\n- Official business name: [NAME]\n- Address/service area: [VERIFIED DETAILS]\n- Contact/CTA: [VERIFIED DETAILS]\n- Services/programmes: [VERIFIED LIST]\n- Target audiences: [VERIFIED LIST]\n- Brand: headings [FONT], body [FONT], colours [HEX VALUES], logo [PATH]\n- Consent: [STATUS FOR PEOPLE/CHILDREN]\n\nDELIVERABLES\n1. Create a responsive, accessible landing page with meaningful HTML, unique metadata and an absolute canonical URL.\n2. Preserve the existing architecture and do not fabricate facts. Keep drafts/private routes out of the sitemap.\n3. Implement a local-first, approval-gated social workflow: detect one eligible approved item; request asset choice; create a deterministic branded preview; validate; require final approval; allow dry-run publishing; block real publishing by default.\n4. Add secure media handling and tests. Never reveal or commit secrets.\n5. Run lint, tests, SEO validation and production build. Report the exact blockers, assumptions and files changed.\n\nSTOP CONDITIONS\n- Do not publish to Facebook, Instagram or any public destination without an explicit final approval for that exact preview.\n- Do not use identifiable children unless written consent is recorded.\n- Do not make paid media-generation calls automatically.\n- Do not schedule unless an approved posting time and time zone exist.")

    heading(doc, "4. Landing-page implementation standard")
    for item in [
        "Start with the verified profile and a page brief: goal, audience, CTA, sections, conversion route and evidence available.",
        "Use semantic headings, descriptive buttons, keyboard-visible focus, image alt text and good contrast. Do not use generic SEO filler.",
        "Set a unique document title, meta description, Open Graph data and absolute canonical URL. Confirm the canonical domain separately from any provisional Vercel URL.",
        "Use only approved, licensed or AI-concept imagery. Label/retain consent data outside public assets. Never portray generated people as real students or instructors.",
        "Keep indexable content server-visible/meaningful in the rendered HTML. Update sitemap only for approved public pages.",
        "Before release, run lint, tests, SEO validation, production build and a browser check. Do not deploy if canonical, metadata or form CTA is unresolved.",
    ]: numbered(doc, item)

    heading(doc, "5. Social workflow and approval state machine")
    table(doc, ["State", "Meaning / permitted next action"], [
        ("detected", "One final-approved, unposted source item is found; no creative or AI call has occurred."),
        ("asset_requested", "Operator must choose: upload image, upload video, repository asset, generate image or skip."),
        ("asset_received / inspected", "Original is preserved, working copy validated; consent, path, MIME type, dimensions and size are recorded."),
        ("creative_ready", "Caption, design manifest and deterministic preview exist. No platform call is permitted."),
        ("awaiting_final_approval", "Preview is shown to the operator. Only a recorded approval may move it forward."),
        ("ready_to_publish", "Validation passed and the exact approved manifest is frozen."),
        ("published / failed / skipped", "Store independent platform result, redacted errors, public URL and idempotency key. Never silently retry duplicate posts."),
    ], [1.7, 4.8])
    add_text(doc, "Required operator sequence", size=12, bold=True, color=GREEN, after=3)
    for cmd in [
        "npm run social:detect",
        "npm run social:request-asset --topic <id>",
        "npm run social:respond-asset --topic <id> --choice <upload-image|upload-video|repository-asset|generate-image|skip>",
        "npm run social:preview --topic <id>",
        "npm run social:approve --topic <id> --action <approve|request_changes|replace_asset|switch_to_image|switch_to_video|skip>",
        "npm run social:publish --topic <id> --dry-run",
    ]: code_block(doc, cmd)
    add_text(doc, "A real publish is a separate, explicit action after the exact preview is approved. The workflow must not create another creative during publishing.", italic=True, color=MUTED, after=8)

    heading(doc, "6. Creative brief and brand layer rules")
    add_text(doc, "Default output: a static 1080 × 1350 post. Create a 1080 × 1080 variant only when needed. A 3–4 second silent video is optional and must use an approved switch; video publishing remains disabled by default.")
    table(doc, ["Element", "Approved production rule"], [
        ("Background", "Authentic or AI-concept martial-arts/fitness visual with subject centre-right or right; retain clear dark/quiet negative space on the left. Avoid aggressive MMA, violence or fake documentary claims."),
        ("Logo", "Use the official logo as a transparent asset at the top-left. Keep roughly 55–70 px top/left margin. Remove any opaque logo background; never place it inside a white box."),
        ("Main line", "A motivating 3–9 word line, maximum two lines, large and wrapped in the left half/lower-left. It must not overlap the subject or cross into the right half."),
        ("Support copy", "One concise, verified supporting line; retain safe readable contrast and spacing."),
        ("CTA", "Place ‘BOOK YOUR FREE TRIAL CLASS’ in the bottom-right with page margins. Use #B85A2B background, #FFFFFF text and modest rounded corners (about 12 px)."),
        ("Contact row", "Instagram handle and verified website at bottom-left. Do not include unverified phone/WhatsApp details."),
        ("Text rendering", "Build readable type as editable SVG/HTML/CSS layers. Do not request generated-image text or rely on it for spelling."),
    ], [1.4, 5.1])
    heading(doc, "7. Caption standard", 2)
    for item in [
        "Generate platform-specific copy from verified configuration and approved source metadata only.",
        "Facebook caption includes the valid blog/landing-page link when one is approved.",
        "Instagram uses ‘link in bio’ only if configuration explicitly confirms that the bio link points to the intended destination.",
        "Use a positive, safe and family-friendly tone. Avoid promises, medical claims, guarantee language or unverified qualifications.",
        "Store caption text and content hash in the creative manifest; publishing must use that frozen approved version.",
    ]: bullet(doc, item)

    heading(doc, "8. Secure Vercel Blob upload endpoint")
    add_text(doc, "Implement a server-side endpoint such as POST /api/social/upload. It is an upload adapter, not a public publishing trigger.")
    table(doc, ["Requirement", "Implementation rule"], [
        ("Authentication", "Require Authorization: Bearer <SOCIAL_UPLOAD_SECRET>. Compare safely; return generic 401/403 errors. Never log the supplied token."),
        ("Blob SDK", "Install and use @vercel/blob from the Vercel server runtime. The connected project must have BLOB_READ_WRITE_TOKEN configured in Vercel environment variables."),
        ("Allowed content", "Accept image/jpeg, image/png, image/webp and video/mp4 only. Enforce content type and byte limits (for example 10 MB image, 50 MB video)."),
        ("Naming", "Store under social/YYYY/MM/DD/<UUID>.<extension>; never trust the client filename. UUID paths are duplicate-safe."),
        ("Response", "Return public Blob URL, pathname, content type and file size only. Do not return credentials or internal configuration."),
        ("Temporary media", "Use a stable public HTTPS URL for Meta. Delete a temporary Blob only after both intended platform results are confirmed; retain it when failure investigation/retry is required."),
    ], [1.45, 5.05])
    code_block(doc, "Required environment variable names (never values):\nSOCIAL_UPLOAD_SECRET\nBLOB_READ_WRITE_TOKEN\nMETA_PAGE_ID\nMETA_INSTAGRAM_ACCOUNT_ID\nMETA_ACCESS_TOKEN\nMETA_API_VERSION\nAUTO_PUBLISH=false\nENABLE_SOCIAL_VIDEO_PUBLISHING=false")

    heading(doc, "9. Meta publishing safeguards")
    for item in [
        "Use a platform client with redacted errors and a dry-run payload summary. Never print access tokens or full API headers.",
        "Before a real image publish, require: approved frozen record, valid public HTTPS media URL, successful public URL check, valid destination URL, non-placeholder business text, correct Meta IDs and required access permissions.",
        "For scheduled posts, additionally require the configured posting time and Asia/Kolkata (or the project’s confirmed) time zone.",
        "Keep video posting behind ENABLE_SOCIAL_VIDEO_PUBLISHING=false until it is deliberately tested and enabled.",
        "Persist Facebook and Instagram results independently. A Facebook success does not imply an Instagram success. Prevent duplicates using topicId:format:platform plus a content/manifest hash.",
        "Do not make Meta calls from GitHub Actions until all required secrets are configured in GitHub and a manual approval input is present.",
    ]: bullet(doc, item)

    heading(doc, "10. GitHub Actions pattern — manual and review-safe")
    table(doc, ["Workflow", "Purpose and safeguards"], [
        ("daily-prepare", "Runs preparation only after a posting time exists. It may detect/review state but never creates a second creative or auto-publishes."),
        ("manual-resume", "Lets an operator resume a paused record from a known approved state. Copies only review-safe JSON to a dedicated review branch."),
        ("manual-publish", "Uses workflow_dispatch with an explicit boolean/string input approve=true. Fails closed for any other value; run dry-run first."),
    ], [1.45, 5.05])
    code_block(doc, "workflow_dispatch:\n  inputs:\n    approve:\n      description: 'Type true only after approving the exact preview'\n      required: true\n      type: string\n\n# Publish job guard\nif: ${{ inputs.approve == 'true' }}")
    add_text(doc, "Do not commit runtime state, generated media, logs or secrets to the normal branch. Mirror only sanitized review artifacts intentionally.", italic=True, color=MUTED)

    heading(doc, "11. Verification checklist")
    for item in [
        "Run unit tests for final-approved detection, stale/pending exclusions, idempotency and state transitions.",
        "Test path traversal, unsupported type, oversize media and identifiable-child consent rejection.",
        "Test placeholder detection, caption link rules and error/secret redaction.",
        "Mock asset intake, mock image-generation pause, preview manifest validation and platform dry-run payloads.",
        "Run the end-to-end fixture: detect → asset decision → inspection → preview → final approval → Facebook/Instagram dry-run → independent persisted results → duplicate prevention.",
        "Run npm lint, tests, SEO validation, social validation, production build and a local image/Remotion render check as applicable.",
        "Do not run paid image generation, deploy, Blob upload or real Meta publishing as part of implementation testing.",
    ]: bullet(doc, item)

    heading(doc, "12. Handoff template for the next implementation")
    code_block(doc, "Project: [NAME]\nGoal: [LANDING PAGE / CAMPAIGN / SOCIAL POST]\nVerified canonical URL: [HTTPS URL]\nApproved CTA and destination: [TEXT + URL]\nVerified offers/audience: [LIST]\nBrand tokens and logo path: [LIST]\nImage source and consent status: [DETAIL]\nSocial accounts/platform IDs: [CONFIGURED OUTSIDE THIS DOCUMENT]\nApproval owner: [NAME/ROLE]\nPost time / time zone: [TIME or 'unset'] / [ZONE]\nAllowed external action today: [NONE | DRY-RUN | UPLOAD | PUBLISH EXACT APPROVED MANIFEST]\nAcceptance checks: [LIST]\nKnown blockers/verification-required fields: [LIST]")
    heading(doc, "Final release decision", 2)
    add_text(doc, "Before a new landing page or post goes live, record: (1) the exact preview/content version, (2) approver and timestamp, (3) verification status of contact and consent data, (4) validation/build results, and (5) platform result IDs and public URLs. If any item is missing, leave the item paused—not published.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
